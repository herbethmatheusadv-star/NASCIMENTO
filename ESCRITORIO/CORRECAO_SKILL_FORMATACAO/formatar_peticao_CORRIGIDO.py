#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VERSAO CORRIGIDA (2026-07-09) do formatador de peticoes Nascimento Advocacia.

RODADA 2 — refinamentos de precisao a partir dos prints do titular (dialogo
"Ajustar Recuos da Lista" do Word) + reextracao via w:numPr/numbering.xml das
3 pecas reais (LUCIANA, EDIO/Aguas do Para, DANIEL).

RODADA 3 — corrigido bug real encontrado ao rodar --docx na peticao real da
DAIANE: pedidos em letra que colidem com algarismos romanos (c, i, l, v, x,
d, m) estavam sendo lidos como titulo; e um fallback de "caixa alta curta"
estava lendo a assinatura final do advogado como titulo novo. Ver notas em
_RE_TITULO_DETECT e no bloco de hints dentro de classificar().

Niveis e valores confirmados:
  - Titulo "I) DOS FATOS": JUSTIFY, negrito, recuo 0/0, "seguir numero com"=espaco.
  - Subtitulo "1.1. Da relacao...": JUSTIFY, negrito+sublinhado, recuo 0/0, espaco.
  - Item numerado "1. 2. 3.": JUSTIFY, recuo primeira linha 2,0cm, numero em
    NEGRITO, texto normal, "seguir numero com"=espaco, numeracao CONTINUA
    (nao reinicia a cada titulo/subtitulo).
  - Pedido "a) b) c)": recuo left=1,27cm hanging=0,635cm, SEM negrito no
    marcador, "seguir numero com"=TAB (default do Word), continuo.
  - Pedido-sub "a. b. c.": recuo left=2,54cm hanging=0,635cm, sem negrito, TAB.
  - Sub-item romano "i. ii. iii." e bullet "•": recuo left=3,81cm
    hanging=0,635cm (o valor cru extraido do XML pro romano era 0,3175cm,
    mas isso estourava com "ii."/"iii." em Century Gothic 11 — ver nota
    em RECUO_SUBITEM_HANGING_CM), TAB.
  - Espacamento: 20pt antes / 10pt depois / linha simples em TODOS os
    paragrafos (titulo, subtitulo, item, pedido, corpo, e ate as linhas em
    branco do bloco antes da qualificacao) — confirmado uniforme nas 3 pecas.
  - Fonte: Century Gothic 11 em tudo.
  - Bloco de espaco em branco antes da qualificacao: 7 a 10 paragrafos em
    branco (varia com o tamanho do texto da qualificacao — nao ha quebra de
    pagina manual, e so espacamento mesmo; ajustar visualmente).
  - Como o marcador e digitado como texto puro (nao e campo de lista nativo
    do Word), o "TAB" dos niveis pedido/sub-item so alinha direito se a
    gente definir uma parada de tabulacao exatamente na posicao do recuo
    (_add_tab_stop) — senao o TAB pula pra um tab-stop generico.

Uso:
  python formatar_peticao_CORRIGIDO.py --texto entrada.md --saida saida.docx
  python formatar_peticao_CORRIGIDO.py --docx entrada.docx --saida saida.docx

Markup do modo texto (linha a linha):
  # Titulo da secao          -> Titulo (I) II) III)... MAIUSCULA, negrito)
  ## Subtitulo               -> Subtitulo (1.1., 1.2., ... negrito+sublinhado)
  - item                     -> paragrafo numerado CONTINUO (1. 2. 3. ...)
                                 com recuo de primeira linha de 2cm
    - item (indentado 2+ espacos) -> sub-item em romano minusculo (i. ii. ...)
    * item (indentado 2+ espacos) -> sub-item em bullet (•) [alternativa]
  + pedido                   -> pedido em letra CONTINUA (a) b) c) ...)
    + pedido (indentado 2+ espacos) -> sub-pedido em letra+ponto (a. b. ...)
  > linha centralizada       -> linha centralizada em negrito
  qualquer outra linha       -> paragrafo de corpo justificado (sem recuo)
  enfase inline: **negrito**  *italico*  __sublinhado__
  linha em branco            -> ignorada (a separacao vem do espacamento)

Parametros:
  --espacos-topo N   numero de paragrafos em branco entre o enderecamento e
                      a qualificacao (padrao 9; qualificacao longa ~1300
                      caracteres -> comece com 7; curta ~700-800 -> comece
                      com 10; ajuste visualmente conforme o caso)
  --acao-centralizada  forca a linha do nome da acao a ficar centralizada
                      (default agora e justificada)
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

FONTE = "Century Gothic"
TAMANHO = 11
ANTES_PT = 20
DEPOIS_PT = 10

# Recuos (extraidos com precisao de w:numPr + numbering.xml das 3 pecas reais —
# cada valor abaixo e uma fracao exata de polegada, que e como o Word realmente
# grava esses recuos por baixo dos paineis em cm; por isso os cm "quebram")
RECUO_ITEM_PRIMEIRA_LINHA_CM = 2.0       # "1. 2. 3." — firstLine=1134 twips (~0,7875in = 2,00cm)
RECUO_PEDIDO_LEFT_CM = 1.27              # "a) b) c)" — left=720 twips (0,5in = 1,27cm exato)
RECUO_PEDIDO_HANGING_CM = 0.635          # hanging=360 twips (0,25in = 0,635cm exato)
RECUO_PEDIDO_SUB_LEFT_CM = 2.54          # "a. b. c." — left=1440 twips (1in = 2,54cm exato)
RECUO_PEDIDO_SUB_HANGING_CM = 0.635      # hanging=360 twips
RECUO_SUBITEM_LEFT_CM = 3.81             # "i. ii. iii." / "•" — left=2160 twips (1,5in = 3,81cm exato)
RECUO_SUBITEM_HANGING_CM = 0.635          # variante romana — AJUSTADO de 0,3175cm (valor
                                          # cru do XML) pra 0,635cm: 0,3175cm (9pt) e
                                          # estreito demais pra "ii." "iii." etc. em
                                          # Century Gothic 11 — o marcador ultrapassa a
                                          # parada de tabulacao e o texto pula pra uma
                                          # parada padrao bem mais longe (bug visual
                                          # confirmado em teste renderizado). 0,635cm
                                          # cabe confortavelmente ate uns 3-4 caracteres.
RECUO_SUBITEM_BULLET_HANGING_CM = 0.635  # variante bullet (•) — hanging=360 twips (peca EDIO)

ESPACOS_TOPO_PADRAO = 9

_ALIGN = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

# --------------------------------------------------------------------------- #
# Helpers de baixo nivel
# --------------------------------------------------------------------------- #

def _set_run_font(run, size_pt=TAMANHO):
    run.font.name = FONTE
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONTE)


def _fmt_para(p, align="justify", antes=ANTES_PT, depois=DEPOIS_PT,
              first_line_cm=0.0, left_cm=0.0, hanging_cm=None):
    """
    first_line_cm: recuo positivo da PRIMEIRA linha (padrao dos itens
                   numerados "1. 2. 3." das pecas reais — a linha do
                   numero entra, as demais voltam pra margem).
    left_cm/hanging_cm: usados para pedidos/sub-itens (recuo tipo "hanging",
                   onde o texto do corpo fica em left_cm e o marcador
                   (a), i., etc.) fica ANTES, em left_cm - hanging_cm.
    """
    pf = p.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(depois)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if hanging_cm is not None:
        pf.left_indent = Cm(left_cm)
        pf.first_line_indent = Cm(-hanging_cm)
    else:
        pf.left_indent = Cm(left_cm) if left_cm else None
        pf.first_line_indent = Cm(first_line_cm) if first_line_cm else Pt(0)


def _add_tab_stop(p, cm_pos):
    """Define uma parada de tabulacao exatamente na posicao do recuo do
    texto, para que um marcador seguido de TAB (ex.: "a)\t", "i.\t", "bullet\t")
    alinhe perfeitamente com as linhas seguintes — do jeito que o Word faz
    nativamente quando "seguir numero com" = tab (padrao real confirmado
    via w:suff em pedido/pedido-sub/sub-item nas 3 pecas)."""
    p.paragraph_format.tab_stops.add_tab_stop(Cm(cm_pos), WD_TAB_ALIGNMENT.LEFT)


def int_to_roman(n, lower=False):
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, s in vals:
        while n >= v:
            out += s
            n -= v
    return out.lower() if lower else out


def int_to_letter(n):
    """1->a, 2->b, ..., 26->z, 27->aa, ..."""
    out = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        out = chr(97 + rem) + out
    return out


def _strip_inline_markers(t):
    return re.sub(r"\*\*|__|\*", "", t)


_RE_TITULO_MARK = re.compile(r"^\s*[IVXLCM]+\s*[\)\.\-–]\s*", re.I)
_RE_SUBT_MARK = re.compile(r"^\s*\d+\.\d+\.?\s*[\)\.\-–]?\s*")
_RE_ITEM_MARK = re.compile(r"^\s*(\d+\s*[\.\)]\s*|[-*•]\s*)")
_RE_PEDIDO_MARK = re.compile(r"^\s*[a-z]\)\s*", re.I)

# Deteccao (nao stripping) de titulo ja digitado: EXIGE maiusculas.
# Sem isso, letras de pedido que tambem sao algarismos romanos validos
# (c, i, l, v, x, d, m) colidiam com titulo — ex.: "c) ordem a re..." e
# "i) a exibicao..." estavam sendo lidos como titulo "C)"/"I)" (bug real
# encontrado ao rodar --docx na peticao da DAIANE: os pedidos em "c)" e
# "i)" da propria peca viravam titulo, saindo em CAIXA ALTA e quebrando a
# numeracao das letras seguintes). Titulo real nesta casa e sempre
# maiusculo ("I)", "II)", "III)", "IV)"); pedido e sempre minusculo — por
# isso a exigencia de maiuscula aqui e segura e resolve a ambiguidade.
_RE_TITULO_DETECT = re.compile(r"^\s*[IVXLCM]+\)\s+\S")


def parse_inline(text):
    token = re.compile(r"(\*\*(?P<b>.+?)\*\*|__(?P<u>.+?)__|\*(?P<i>.+?)\*)")
    out, last = [], 0
    for m in token.finditer(text):
        if m.start() > last:
            out.append((text[last:m.start()], False, False, False))
        if m.group("b") is not None:
            out.append((m.group("b"), True, False, False))
        elif m.group("u") is not None:
            out.append((m.group("u"), False, False, True))
        elif m.group("i") is not None:
            out.append((m.group("i"), False, True, False))
        last = m.end()
    if last < len(text):
        out.append((text[last:], False, False, False))
    return out or [(text, False, False, False)]


def _add_runs(p, texto):
    for seg, b, i, u in parse_inline(texto):
        r = p.add_run(seg)
        _set_run_font(r)
        if b:
            r.bold = True
        if i:
            r.italic = True
        if u:
            r.underline = True


# --------------------------------------------------------------------------- #
# Construtores de paragrafo por nivel
# --------------------------------------------------------------------------- #

def add_titulo(doc, texto, idx):
    texto = _strip_inline_markers(_RE_TITULO_MARK.sub("", texto.strip())).strip().upper()
    p = doc.add_paragraph()
    _fmt_para(p, align="justify")  # confirmado: titulo e JUSTIFY nas 3 pecas, nao LEFT
    run = p.add_run(f"{int_to_roman(idx)}) {texto}")
    _set_run_font(run)
    run.bold = True
    return p


def add_subtitulo(doc, texto, ti, si):
    texto = _strip_inline_markers(_RE_SUBT_MARK.sub("", texto.strip())).strip()
    p = doc.add_paragraph()
    _fmt_para(p, align="justify")
    run = p.add_run(f"{ti}.{si}. {texto}")
    _set_run_font(run)
    run.bold = True
    run.underline = True
    return p


def add_item(doc, texto, n):
    """Paragrafo narrativo numerado "1. 2. 3." — CONTINUO, com recuo de
    primeira linha de 2cm (padrao real; o script anterior nao tinha isso)."""
    texto = _RE_ITEM_MARK.sub("", texto.strip())
    p = doc.add_paragraph()
    _fmt_para(p, align="justify", first_line_cm=RECUO_ITEM_PRIMEIRA_LINHA_CM)
    pre = p.add_run(f"{n}. ")
    _set_run_font(pre)
    pre.bold = True  # confirmado no rPr do nivel (EDIO/DANIEL): o numero do item e negrito
    _add_runs(p, texto)
    return p


def add_subitem_romano(doc, texto, n):
    """"i. ii. iii." — itens dentro de um fato numerado (ex.: lista de
    dividas, faturas). Contador reinicia a cada fato-pai (novo_item() zera
    o contador; ver Numerador). Usa TAB (nao espaco) apos o marcador, com
    parada de tabulacao na propria posicao do recuo."""
    texto = re.sub(r"^\s*[-*•]\s*", "", texto.strip())
    p = doc.add_paragraph()
    _fmt_para(p, align="justify", left_cm=RECUO_SUBITEM_LEFT_CM,
              hanging_cm=RECUO_SUBITEM_HANGING_CM)
    _add_tab_stop(p, RECUO_SUBITEM_LEFT_CM)
    pre = p.add_run(f"{int_to_roman(n, lower=True)}.\t")
    _set_run_font(pre)
    _add_runs(p, texto)
    return p


def add_subitem_bullet(doc, texto):
    """Alternativa em bullet, como em uma das pecas de referencia.
    Usa o mesmo recuo pendurado da variante romana (0,635cm) — o valor
    cru do XML era menor (0,3175cm) mas estourava com marcadores de 2+
    caracteres; ver nota em RECUO_SUBITEM_HANGING_CM."""
    texto = re.sub(r"^\s*[-*•]\s*", "", texto.strip())
    p = doc.add_paragraph()
    _fmt_para(p, align="justify", left_cm=RECUO_SUBITEM_LEFT_CM,
              hanging_cm=RECUO_SUBITEM_BULLET_HANGING_CM)
    _add_tab_stop(p, RECUO_SUBITEM_LEFT_CM)
    pre = p.add_run("•\t")
    _set_run_font(pre)
    _add_runs(p, texto)
    return p


def add_pedido(doc, texto, n):
    """"a) b) c)..." CONTINUO — usado em tutela de urgencia / dos pedidos.
    O marcador NAO e negrito (confirmado no rPr do nivel — diferente do
    item numerado, que e negrito). Usa TAB apos o marcador."""
    texto = _RE_PEDIDO_MARK.sub("", texto.strip())
    p = doc.add_paragraph()
    _fmt_para(p, align="justify", left_cm=RECUO_PEDIDO_LEFT_CM,
              hanging_cm=RECUO_PEDIDO_HANGING_CM)
    _add_tab_stop(p, RECUO_PEDIDO_LEFT_CM)
    pre = p.add_run(f"{int_to_letter(n)})\t")
    _set_run_font(pre)
    _add_runs(p, texto)
    return p


def add_pedido_sub(doc, texto, n):
    """"a. b. c." — sub-itens de um pedido. Tambem nao-negrito e com TAB."""
    texto = re.sub(r"^\s*[a-z][\.\)]\s*", "", texto.strip(), flags=re.I)
    p = doc.add_paragraph()
    _fmt_para(p, align="justify", left_cm=RECUO_PEDIDO_SUB_LEFT_CM,
              hanging_cm=RECUO_PEDIDO_SUB_HANGING_CM)
    _add_tab_stop(p, RECUO_PEDIDO_SUB_LEFT_CM)
    pre = p.add_run(f"{int_to_letter(n)}.\t")
    _set_run_font(pre)
    _add_runs(p, texto)
    return p


def add_corpo(doc, texto, align="justify", negrito_base=False):
    p = doc.add_paragraph()
    _fmt_para(p, align=align)
    for seg, b, i, u in parse_inline(texto):
        r = p.add_run(seg)
        _set_run_font(r)
        if b or negrito_base:
            r.bold = True
        if i:
            r.italic = True
        if u:
            r.underline = True
    return p


def add_centralizado(doc, texto):
    texto = _strip_inline_markers(texto.strip())
    p = doc.add_paragraph()
    _fmt_para(p, align="center")
    run = p.add_run(texto)
    _set_run_font(run)
    run.bold = True
    return p


def add_blank(doc):
    """Paragrafo em branco com o mesmo espacamento — usado no bloco de
    espaco antes da qualificacao (padrao real: 7 a 10 paragrafos, ajustar
    conforme o tamanho da qualificacao)."""
    p = doc.add_paragraph()
    _fmt_para(p, align="justify")
    return p


# --------------------------------------------------------------------------- #
# Preparo do documento base (timbrado)
# --------------------------------------------------------------------------- #

def _limpar_corpo(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _aplicar_estilo_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = FONTE
    style.font.size = Pt(TAMANHO)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONTE)


def _aplicar_margens(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)


def preparar_doc(template_path):
    if not os.path.exists(template_path):
        sys.exit(f"ERRO: timbrado nao encontrado: {template_path}")
    doc = Document(template_path)
    _limpar_corpo(doc)
    _aplicar_estilo_normal(doc)
    _aplicar_margens(doc)
    return doc


# --------------------------------------------------------------------------- #
# Classificador de linhas
# --------------------------------------------------------------------------- #

def _indent_level(linha_bruta):
    """Conta espacos/tabs no inicio pra distinguir item de sub-item."""
    m = re.match(r"^(\s*)", linha_bruta)
    return len(m.group(1).replace("\t", "  "))


def classificar(texto, hints=None):
    t = texto.rstrip()
    if not t.strip():
        return ("blank", "", 0)

    indent = _indent_level(texto)

    if t.lstrip().startswith("## "):
        return ("subtitulo", t.lstrip()[3:].strip(), 0)
    if t.lstrip().startswith("# "):
        return ("titulo", t.lstrip()[2:].strip(), 0)
    if t.lstrip().startswith("> "):
        return ("center", t.lstrip()[2:].strip(), 0)

    # pedidos: "+" no inicio (nivel 0) ou indentado (nivel 1 = sub-pedido)
    if re.match(r"^\s*\+\s+\S", t):
        nivel = "pedido-sub" if indent >= 2 else "pedido"
        return (nivel, re.sub(r"^\s*\+\s+", "", t), 0)

    # itens numerados "-" (nivel 0) vs sub-itens indentados (nivel 1+)
    if re.match(r"^\s*[-*]\s+\S", t):
        if indent >= 2:
            marcador = t.lstrip()[0]
            nivel = "subitem-bullet" if marcador == "*" else "subitem-romano"
            return (nivel, re.sub(r"^\s*[-*]\s+", "", t), 0)
        return ("item", re.sub(r"^\s*[-*]\s+", "", t), 0)
    if re.match(r"^\s*•\s+\S", t):
        return ("subitem-bullet", re.sub(r"^\s*•\s+", "", t), 0)

    # marcadores ja digitados (texto colado / modo docx)
    if _RE_TITULO_DETECT.match(t):
        return ("titulo", _RE_TITULO_MARK.sub("", t), 0)
    if re.match(r"^\s*\d+\.\d+\.?\s+\S", t):
        return ("subtitulo", _RE_SUBT_MARK.sub("", t), 0)
    # item numerado "1. Texto..." (checado DEPOIS do subtitulo "1.1." pra nao colidir)
    if re.match(r"^\s*\d+\.\s+\S", t):
        return ("item", _RE_ITEM_MARK.sub("", t, count=1), 0)
    if re.match(r"^\s*[a-z]\)\s+\S", t, re.I):
        return ("pedido", _RE_PEDIDO_MARK.sub("", t), 0)

    if hints:
        # OBS: removido o fallback "CAIXA ALTA curta e sem ponto final -> titulo".
        # Causava falso positivo real: a assinatura final do advogado
        # ("CARLOS EDUARDO DE SOUSA DO NASCIMENTO", em negrito e caixa alta,
        # sem ponto final) estava sendo lida como um titulo novo. Todos os
        # titulos reais desta casa ja vem com marcador explicito ("I)",
        # "II)"...), capturado antes de chegar aqui — este fallback nao e
        # necessario para o padrao real e so cria ambiguidade com nomes/
        # assinaturas em negrito.
        if hints.get("bold") and hints.get("underline") and len(t.strip()) <= 130:
            return ("subtitulo", t.strip(), 0)
        if hints.get("align") == "center":
            return ("center", t.strip(), 0)

    return ("corpo", t.strip(), 0)


class Numerador:
    """Titulo/subtitulo por hierarquia; item e pedido CONTINUOS (nao
    reiniciam a cada subtitulo/titulo — padrao real conferido nas 3 pecas).
    Sub-item (romano) reinicia a cada item-pai novo."""

    def __init__(self):
        self.titulo = 0
        self.subtitulo = 0
        self.item = 0
        self.pedido = 0
        self.pedido_sub = 0
        self.subitem = 0

    def novo_titulo(self):
        self.titulo += 1
        self.subtitulo = 0
        return self.titulo

    def novo_subtitulo(self):
        if self.titulo == 0:
            self.titulo = 1
        self.subtitulo += 1
        return self.titulo, self.subtitulo

    def novo_item(self):
        self.item += 1
        self.subitem = 0
        return self.item

    def novo_subitem(self):
        self.subitem += 1
        return self.subitem

    def novo_pedido(self):
        self.pedido += 1
        self.pedido_sub = 0
        return self.pedido

    def novo_pedido_sub(self):
        self.pedido_sub += 1
        return self.pedido_sub


# --------------------------------------------------------------------------- #
# Modo A — a partir de texto/markup
# --------------------------------------------------------------------------- #

def gerar_de_texto(texto, template_path, saida, negrito_corpo=False,
                    espacos_topo=ESPACOS_TOPO_PADRAO, acao_centralizada=False):
    doc = preparar_doc(template_path)
    num = Numerador()

    linhas = texto.splitlines()

    primeira_nao_vazia = next((i for i, l in enumerate(linhas) if l.strip()), None)

    for idx, linha in enumerate(linhas):
        nivel, limpo, _ = classificar(linha)

        if nivel == "blank":
            continue

        if nivel == "titulo":
            add_titulo(doc, limpo, num.novo_titulo())
        elif nivel == "subtitulo":
            ti, si = num.novo_subtitulo()
            add_subtitulo(doc, limpo, ti, si)
        elif nivel == "item":
            add_item(doc, limpo, num.novo_item())
        elif nivel == "subitem-romano":
            add_subitem_romano(doc, limpo, num.novo_subitem())
        elif nivel == "subitem-bullet":
            add_subitem_bullet(doc, limpo)
        elif nivel == "pedido":
            add_pedido(doc, limpo, num.novo_pedido())
        elif nivel == "pedido-sub":
            add_pedido_sub(doc, limpo, num.novo_pedido_sub())
        elif nivel == "center":
            add_centralizado(doc, limpo)
        else:
            align = "center" if acao_centralizada and _parece_nome_acao(limpo) else "justify"
            add_corpo(doc, limpo, align=align, negrito_base=negrito_corpo)

        if idx == primeira_nao_vazia:
            for _ in range(espacos_topo):
                add_blank(doc)

    doc.save(saida)
    return saida


def _parece_nome_acao(t):
    return t.upper().startswith(("AÇÃO", "ACAO"))


# --------------------------------------------------------------------------- #
# Modo B — reformatando um .docx existente
# --------------------------------------------------------------------------- #

def _hints_do_paragrafo(p):
    runs = [r for r in p.runs if (r.text or "").strip()]
    if not runs:
        runs = p.runs

    def maioria(attr):
        vals = [getattr(r, attr) for r in runs]
        vals = [v for v in vals if v is not None]
        return bool(vals) and sum(1 for v in vals if v) >= (len(vals) / 2)

    align = {WD_ALIGN_PARAGRAPH.CENTER: "center",
             WD_ALIGN_PARAGRAPH.RIGHT: "right"}.get(p.alignment, None)
    return {"bold": maioria("bold"), "underline": maioria("underline"), "align": align}


def gerar_de_docx(origem, template_path, saida, corpo="preservar",
                   espacos_topo=ESPACOS_TOPO_PADRAO):
    if not os.path.exists(origem):
        sys.exit(f"ERRO: arquivo de origem nao encontrado: {origem}")
    src = Document(origem)
    doc = preparar_doc(template_path)
    num = Numerador()

    paras = src.paragraphs
    primeira_nao_vazia = next((i for i, p in enumerate(paras) if p.text.strip()), None)

    for i, p in enumerate(paras):
        texto = p.text
        if not texto.strip():
            continue
        hints = _hints_do_paragrafo(p)
        nivel, limpo, _ = classificar(texto, hints=hints)

        if nivel == "titulo":
            add_titulo(doc, limpo, num.novo_titulo())
        elif nivel == "subtitulo":
            ti, si = num.novo_subtitulo()
            add_subtitulo(doc, limpo, ti, si)
        elif nivel == "item":
            n = num.novo_item()
            np = doc.add_paragraph()
            _fmt_para(np, align="justify", first_line_cm=RECUO_ITEM_PRIMEIRA_LINHA_CM)
            pre = np.add_run(f"{n}. ")
            _set_run_font(pre)
            pre.bold = True
            _copiar_runs_stripping_marcador(p, np, corpo, _RE_ITEM_MARK)
        elif nivel == "pedido":
            np = doc.add_paragraph()
            n = num.novo_pedido()
            _fmt_para(np, align="justify", left_cm=RECUO_PEDIDO_LEFT_CM,
                      hanging_cm=RECUO_PEDIDO_HANGING_CM)
            _add_tab_stop(np, RECUO_PEDIDO_LEFT_CM)
            pre = np.add_run(f"{int_to_letter(n)})\t")
            _set_run_font(pre)
            _copiar_runs_stripping_marcador(p, np, corpo, _RE_PEDIDO_MARK)
        elif nivel == "center":
            np = doc.add_paragraph()
            _fmt_para(np, align="center")
            _copiar_runs(p, np, corpo)
        else:
            align = "justify"
            if hints.get("align") == "right":
                align = "right"
            np = doc.add_paragraph()
            _fmt_para(np, align=align)
            _copiar_runs(p, np, corpo)

        if i == primeira_nao_vazia:
            for _ in range(espacos_topo):
                add_blank(doc)

    doc.save(saida)
    return saida


def _copiar_runs(src_p, dst_p, corpo):
    runs = src_p.runs
    if not runs:
        r = dst_p.add_run(src_p.text)
        _set_run_font(r)
        return
    for r in runs:
        nr = dst_p.add_run(r.text)
        _set_run_font(nr)
        if r.italic:
            nr.italic = True
        if r.underline:
            nr.underline = True
        if r.bold and corpo != "normal":
            nr.bold = True


def _copiar_runs_stripping_marcador(src_p, dst_p, corpo, mark_re):
    """Copia os runs do paragrafo de origem, removendo o marcador (numero,
    letra, bullet etc.) do INICIO do primeiro run nao-vazio.

    Necessario porque, ao reformatar um .docx ja existente (--docx), o
    marcador ("1.", "a)") normalmente ja vem como texto literal dentro do
    proprio paragrafo de origem (as vezes ate como um run isolado, em
    negrito). Sem isso, o marcador novo que a funcao chamadora acabou de
    adicionar (ex.: "1. ", "a)\t") ficaria DUPLICADO com o que ja existe
    no paragrafo copiado (ex.: "1. 1. Texto..."). Se um run ficar vazio
    apos a remocao (era só o marcador), ele e pulado inteiramente.
    """
    runs = list(src_p.runs)
    if not runs:
        texto = mark_re.sub("", src_p.text, count=1)
        r = dst_p.add_run(texto)
        _set_run_font(r)
        return

    ja_removeu = False
    for r in runs:
        texto = r.text
        if not ja_removeu:
            novo_texto = mark_re.sub("", texto, count=1)
            if novo_texto != texto:
                ja_removeu = True
            texto = novo_texto
        if texto == "":
            continue
        nr = dst_p.add_run(texto)
        _set_run_font(nr)
        if r.italic:
            nr.italic = True
        if r.underline:
            nr.underline = True
        if r.bold and corpo != "normal":
            nr.bold = True


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #

def _default_template():
    here = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(here, "..", "assets", "Timbrado Oficial Nascimento.docx")


def main():
    ap = argparse.ArgumentParser(
        description="Aplica o padrao de formatacao de peticoes (timbrado Nascimento) - versao corrigida.")
    src = ap.add_mutually_exclusive_group(required=True)
    src.add_argument("--texto", help="arquivo .md/.txt com o conteudo/markup")
    src.add_argument("--texto-stdin", action="store_true")
    src.add_argument("--docx", help="arquivo .docx existente para reformatar")
    ap.add_argument("--saida", required=True)
    ap.add_argument("--template", default=None)
    ap.add_argument("--corpo", choices=["normal", "preservar"], default=None)
    ap.add_argument("--espacos-topo", type=int, default=ESPACOS_TOPO_PADRAO)
    ap.add_argument("--acao-centralizada", action="store_true")
    args = ap.parse_args()

    template = args.template or _default_template()

    if args.docx:
        corpo = args.corpo or "preservar"
        out = gerar_de_docx(args.docx, template, args.saida, corpo=corpo,
                             espacos_topo=args.espacos_topo)
    else:
        if args.texto_stdin:
            conteudo = sys.stdin.read()
        else:
            with open(args.texto, encoding="utf-8") as fh:
                conteudo = fh.read()
        negrito_corpo = (args.corpo == "preservar")
        out = gerar_de_texto(conteudo, template, args.saida, negrito_corpo=negrito_corpo,
                              espacos_topo=args.espacos_topo,
                              acao_centralizada=args.acao_centralizada)

    print(f"OK -> {out}")


if __name__ == "__main__":
    main()
