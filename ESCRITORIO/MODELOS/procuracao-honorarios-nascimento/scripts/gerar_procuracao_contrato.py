#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera PROCURAÇÃO (ad judicia et extra) e/ou CONTRATO DE HONORÁRIOS (padrão Juizado
Especial Cível) do escritório NASCIMENTO ADVOGADOS, dentro do timbrado oficial.

Só muda o que varia por caso: a QUALIFICAÇÃO do cliente, o OBJETO (nome da ação) e
o(s) RÉU(S). Todo o resto (poderes da procuração, cláusulas 2ª–11ª do contrato,
contratado/advogado, foro) é padrão e já vem embutido. Saída SEMPRE em .docx.

Uso:
    python gerar_procuracao_contrato.py --dados dados.json [--saida-dir PASTA]
                                        [--data "22 de julho de 2026"]
                                        [--so procuracao|contrato|ambos]

dados.json (campos):
    cliente        (obrigatório) — qualificação COMPLETA do cliente, começando pelo
                   NOME em CAIXA ALTA. Ex.:
                   "EDSON RIBEIRO DIAS, brasileiro, solteiro, portador da Cédula de
                    Identidade RG n.º 424030 PC/PA, inscrito no CPF/MF sob o n.º
                    895.555.542-34, residente e domiciliado à Rua H, n.º 21, Bairro
                    União, Parauapebas/PA, CEP 68.515-000"
    cliente_nome   (opcional) — nome p/ assinatura e nome do arquivo. Se ausente, é
                   deduzido do trecho antes da primeira vírgula de "cliente".
    cliente_cpf    (opcional) — CPF exibido sob a assinatura (senão tenta extrair).
    objeto         (contrato) — nome da ação. Ex.: "AÇÃO DE RESCISÃO CONTRATUAL POR
                   VÍCIO DE CONSENTIMENTO C/C RESTITUIÇÃO INTEGRAL DE VALORES,
                   INDENIZAÇÃO POR DANOS MORAIS"
    reus           (contrato) — lista de strings, cada uma a qualificação de um réu.
                   Envolva o NOME da empresa em **negrito**. Ex.:
                   ["**J. FERREIRA REPRESENTAÇÕES LTDA**, pessoa jurídica de direito
                     privado, inscrita no CNPJ/ME nº 43.674.644/0001-78, situada na
                     Avenida Rio Grande, nº 155, Bairro Beiro Rio, Parauapebas – PA,
                     CEP 68515-000"]
    honorarios     (opcional) — percentual de êxito. Padrão: "50% (CINQUENTA POR CENTO)".
    data           (opcional) — texto da data. Padrão: "Parauapebas/PA, <hoje>."

Observação: para procuração basta "cliente". objeto/reus só são exigidos p/ o contrato.
"""
import argparse, json, os, re, sys, datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONTE, TAM, ANTES, DEPOIS = "Century Gothic", 11, 20, 10
_ALIGN = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
          "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
MESES = ["janeiro","fevereiro","março","abril","maio","junho","julho","agosto",
         "setembro","outubro","novembro","dezembro"]

# ---- dados FIXOS do escritório (advogado/contratado) ----
OUTORGADO = ("HERBETH MATHEUS MENDONÇA DO NASCIMENTO, brasileiro, advogado, inscrito na "
    "OAB/PA sob o n.º 39.261, com escritório profissional na Rua D, n.º 410, Sala B, "
    "Bairro Cidade Nova, Parauapebas/PA, CEP 68.310-418, endereço eletrônico "
    "herbethmatheus.adv@gmail.com, onde recebe intimações e notificações")
CONTRATADO = ("HERBETH MATHEUS MENDONÇA DO NASCIMENTO, brasileiro, solteiro, advogado, "
    "inscrito na OAB/PA sob o n.º 39.261, com escritório profissional na Rua D, n.º 410, "
    "Sala B, Bairro Cidade Nova, Parauapebas/PA, CEP 68.310-418, e-mail: "
    "herbethmatheus.adv@gmail.com")
ADV_ASSINA = ("HERBETH MATHEUS MENDONÇA DO NASCIMENTO", "Contratado — OAB/PA nº 39.261")

def _default_template():
    here = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(here, "..", "assets", "Timbrado Oficial Nascimento.docx")

# ---------- motor de formatação (padrão Nascimento) ----------
def _set_font(run, size=TAM):
    run.font.name = FONTE; run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a), FONTE)

def _fmt(p, align="justify", antes=ANTES, depois=DEPOIS):
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(antes), Pt(depois)
    pf.line_spacing = 1.0; pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Pt(0); pf.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

def _inline(text):
    tok = re.compile(r"(\*\*(?P<b>.+?)\*\*|__(?P<u>.+?)__|\*(?P<i>.+?)\*)")
    out, last = [], 0
    for m in tok.finditer(text):
        if m.start() > last: out.append((text[last:m.start()], False, False, False))
        if m.group("b") is not None: out.append((m.group("b"), True, False, False))
        elif m.group("u") is not None: out.append((m.group("u"), False, False, True))
        elif m.group("i") is not None: out.append((m.group("i"), False, True, False))
        last = m.end()
    if last < len(text): out.append((text[last:], False, False, False))
    return out or [(text, False, False, False)]

def par(doc, text, align="justify", bold=False, underline=False, italic=False,
        antes=ANTES, depois=DEPOIS, size=TAM):
    p = doc.add_paragraph(); _fmt(p, align, antes, depois)
    for seg, b, i, u in _inline(text):
        r = p.add_run(seg); _set_font(r, size)
        if bold or b: r.bold = True
        if italic or i: r.italic = True
        if underline or u: r.underline = True
    return p

def heading(doc, text):
    return par(doc, text, align="left", bold=True, underline=True)

def linha_assinatura(doc, largura_cm=8.0):
    ind = (17.0 - largura_cm) / 2.0  # A4 útil 17cm
    p = doc.add_paragraph(); _fmt(p, "center", 0, 0)
    p.paragraph_format.left_indent = Cm(ind); p.paragraph_format.right_indent = Cm(ind)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6"); b.set(qn("w:space"),"1"); b.set(qn("w:color"),"000000")
    pbdr.append(b); pPr.append(pbdr)
    _set_font(p.add_run(" "))
    return p

def assinatura(doc, nome, papel):
    par(doc, "", antes=40, depois=0)
    linha_assinatura(doc)
    par(doc, f"**{nome}**", align="center", antes=3, depois=0)
    par(doc, papel, align="center", antes=0, depois=6)

def preparar(template):
    if not os.path.exists(template): sys.exit(f"ERRO: timbrado não encontrado: {template}")
    doc = Document(template)
    body = doc.element.body
    for c in list(body):
        if c.tag != qn("w:sectPr"): body.remove(c)
    st = doc.styles["Normal"]; st.font.name = FONTE; st.font.size = Pt(TAM)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a), FONTE)
    s = doc.sections[0]; s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(2), Cm(2)
    return doc

# ---------- documentos ----------
def gerar_procuracao(d, template, saida_dir, data_txt):
    doc = preparar(template)
    par(doc, "PROCURAÇÃO", align="center", bold=True, size=15, antes=6, depois=2)
    par(doc, "*Ad Judicia et Extra — Poderes Gerais e Especiais*", align="center", antes=0, depois=14)
    par(doc, f"**OUTORGANTE:** {d['cliente']}.")
    par(doc, f"**OUTORGADO:** {OUTORGADO}.")
    par(doc,
        "**PODERES:** Pelo presente instrumento particular, o **OUTORGANTE** nomeia e "
        "constitui o **OUTORGADO** como seu bastante procurador, conferindo-lhe amplos "
        "poderes para o foro em geral, com a cláusula “ad judicia et extra” (art. 105 do "
        "CPC), em qualquer Juízo, Instância ou Tribunal, podendo propor contra quem de "
        "direito as ações competentes e defendê-lo nas contrárias, seguindo umas e outras "
        "até final decisão, usando os recursos legais e acompanhando-os, conferindo-lhe "
        "ainda poderes especiais para: confessar, reconhecer a procedência do pedido, "
        "transigir, desistir, renunciar ao direito sobre o qual se funda a ação, receber, "
        "dar quitação, firmar acordos e compromissos, levantar alvarás e valores "
        "depositados em juízo, requerer gratuidade da justiça e firmar declaração de "
        "hipossuficiência econômica, prestar depoimento pessoal, interpor recursos "
        "(inclusive agravo de instrumento, apelação, embargos de declaração, recurso "
        "especial e extraordinário), opor exceções e impugnações, requerer tutelas de "
        "urgência e antecipadas, peticionar perante os sistemas SISBAJUD, RENAJUD, INFOJUD, "
        "SERASAJUD, CNIB, CCS-BACEN, requerer desbloqueio de valores e cancelamento de "
        "ordens de penhora online, atuar perante órgãos públicos federais, estaduais e "
        "municipais, INSS, Receita Federal, Detran, cartórios e tabelionatos, podendo "
        "substabelecer esta a outrem, com ou sem reserva de iguais poderes, para agir em "
        "conjunto ou separadamente com o substabelecido, dando tudo por bom, firme e valioso.")
    par(doc, data_txt, align="right", antes=18, depois=6)
    assinatura(doc, d["cliente_nome"], f"Outorgante — CPF n.º {d['cliente_cpf']}")
    out = os.path.join(saida_dir, f"PROCURAÇÃO - {d['cliente_nome']}.docx")
    doc.save(out); return out

def gerar_contrato(d, template, saida_dir, data_txt):
    doc = preparar(template)
    par(doc, "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS",
        align="center", bold=True, underline=True, size=13, antes=6, depois=12)
    par(doc,
        "Pelo presente instrumento particular, as partes abaixo qualificadas têm entre si "
        "justo e contratado o presente **Contrato de Prestação de Serviços Advocatícios**, "
        "regido pelo Estatuto da Advocacia e da OAB (Lei nº 8.906/94), pelo Código de Ética "
        "e Disciplina da OAB, pelo Código de Processo Civil (Lei nº 13.105/2015), pela Lei "
        "dos Juizados Especiais (Lei nº 9.099/95) e pelo Código Civil (Lei nº 10.406/2002), "
        "mediante as cláusulas e condições a seguir:")
    par(doc, f"**CONTRATANTE:** {d['cliente']}.")
    par(doc, f"**CONTRATADO:** {CONTRATADO}.")

    heading(doc, "CLÁUSULA 1ª — DO OBJETO")
    par(doc,
        "O **CONTRATADO**, mediante o competente mandado judicial outorgado pelo "
        "**CONTRATANTE**, obriga-se a prestar serviços advocatícios consistentes na "
        f"propositura e acompanhamento de **{d['objeto']}**, a ser distribuída em face de:")
    for i, reu in enumerate(d["reus"], 1):
        par(doc, f"{i}. {reu}." if not reu.rstrip().endswith((".",";")) else f"{i}. {reu}")
    par(doc,
        "**Parágrafo Único.** Os serviços abrangem o patrocínio da causa em primeira "
        "instância (Juizado Especial Cível da Comarca de Parauapebas/PA) e em segunda "
        "instância (Turma Recursal), incluindo a fase de cumprimento de sentença/execução, "
        "até o efetivo recebimento dos valores pelo **CONTRATANTE**. A atuação em instâncias "
        "superiores (STJ e STF) dependerá de ajuste complementar específico.")

    heading(doc, "CLÁUSULA 2ª — DAS OBRIGAÇÕES DO CONTRATADO")
    par(doc, "São obrigações do **CONTRATADO**:")
    par(doc, "**a)** Prestar os serviços advocatícios com zelo, lealdade, ética e nos prazos legais;")
    par(doc, "**b)** Manter o **CONTRATANTE** informado sobre o andamento processual relevante, preferencialmente por meio eletrônico (WhatsApp);")
    par(doc, "**c)** Guardar sigilo profissional sobre todos os fatos e documentos;")
    par(doc, "**d)** Conservar a documentação fornecida e devolvê-la ao término dos serviços, se solicitado.")

    heading(doc, "CLÁUSULA 3ª — DAS OBRIGAÇÕES DO CONTRATANTE")
    par(doc, "São obrigações do **CONTRATANTE**:")
    par(doc, "**a)** Fornecer ao **CONTRATADO** todos os documentos, informações e provas necessárias ao bom desempenho do mandato, declarando sob as penas da lei a veracidade dos fatos narrados;")
    par(doc, "**b)** Comparecer a **TODAS** as audiências para as quais for designado, sob pena de extinção do processo e condenação ao pagamento de custas processuais (Art. 51, inciso I, da Lei 9.099/95), ônus que recairá exclusivamente sobre o **CONTRATANTE**;")
    par(doc, "**c)** Comunicar imediatamente qualquer alteração de endereço, telefone ou e-mail;")
    par(doc, "**d)** Não realizar acordo, transação ou qualquer composição diretamente com a parte contrária sem a prévia e expressa anuência do **CONTRATADO**, sob pena de incidência da cláusula penal prevista neste instrumento.")

    heading(doc, "CLÁUSULA 4ª — DOS HONORÁRIOS CONTRATUAIS DE ÊXITO")
    par(doc,
        "Em remuneração aos serviços profissionais ora pactuados, o **CONTRATANTE** pagará "
        f"ao **CONTRATADO**, a título de honorários contratuais de êxito, o percentual de "
        f"**{d['honorarios']}** incidente sobre o **VALOR BRUTO** do proveito econômico "
        "obtido ao final da demanda, seja por sentença transitada em julgado, tutela de "
        "urgência, ou acordo judicial/extrajudicial.")
    par(doc,
        "**Parágrafo Primeiro:** Compreende-se por proveito econômico bruto a totalidade dos "
        "valores recebidos, incluindo: restituição de valores (simples ou em dobro), "
        "indenização por danos morais, danos materiais, juros, correção monetária e "
        "eventuais multas cominatórias (astreintes) fixadas por descumprimento de ordem judicial.")
    par(doc,
        "**Parágrafo Segundo (Cláusula de Retenção):** Fica o **CONTRATADO** expressamente "
        "autorizado a reter e compensar os honorários contratuais devidos diretamente dos "
        "valores levantados através de alvará judicial, requisição de pequeno valor (RPV) ou "
        "depósitos realizados pelas rés, efetuando o repasse do saldo remanescente ao **CONTRATANTE**.")

    heading(doc, "CLÁUSULA 5ª — DOS HONORÁRIOS DE SUCUMBÊNCIA")
    par(doc,
        "Os honorários advocatícios de sucumbência, fixados judicialmente nos termos do art. "
        "85 do CPC e art. 55 da Lei 9.099/95, pertencem **EXCLUSIVAMENTE** ao **CONTRATADO**, "
        "possuem natureza alimentar e autônoma, e **NÃO** se confundem nem se compensam com "
        "os honorários contratuais previstos na Cláusula 4ª.")

    heading(doc, "CLÁUSULA 6ª — DAS DESPESAS PROCESSUAIS E JUSTIÇA GRATUITA")
    par(doc, "O acesso ao Juizado Especial Cível em primeira instância independe do pagamento de custas, taxas ou despesas (Art. 54 da Lei 9.099/95).")
    par(doc,
        "**Parágrafo Único:** Caso haja necessidade de interposição de recurso ou o feito "
        "seja remetido à Justiça Comum, o **CONTRATADO** requererá os benefícios da "
        "**Justiça Gratuita**. Caso o benefício seja negado pelo juiz, o **CONTRATANTE** "
        "arcará com as custas processuais (preparo) estritamente necessárias para o andamento do feito.")

    heading(doc, "CLÁUSULA 7ª — DOS VALORES ANTECIPADOS (TUTELAS/LIMINARES)")
    par(doc,
        "Caso o **CONTRATANTE** venha a obter o objeto deste contrato administrativamente "
        "após a contratação, ou receba valores antecipadamente por força de liminar/tutela "
        "de urgência, incidirão imediatamente os honorários contratuais previstos na Cláusula "
        "4ª, devendo o repasse/pagamento ocorrer no prazo máximo de 48 (quarenta e oito) "
        "horas do recebimento.")

    heading(doc, "CLÁUSULA 8ª — DA DESISTÊNCIA, RENÚNCIA OU ACORDO SEM ANUÊNCIA")
    par(doc,
        "Caso o **CONTRATANTE**, sem prévia anuência por escrito do **CONTRATADO**: (a) "
        "desista da ação/procedimento; (b) realize acordo direto com a parte contrária; ou "
        "(c) pratique ato que cause a extinção do processo sem julgamento do mérito (como "
        "faltar a audiências); o contrato será considerado rescindido por culpa do cliente, "
        "vencendo-se antecipadamente os honorários advocatícios contratuais na integralidade, "
        "calculados sobre o valor total do pedido ou do acordo (o que for maior).")

    heading(doc, "CLÁUSULA 9ª — DA RESCISÃO E REVOGAÇÃO DO MANDATO")
    par(doc,
        "A revogação do mandato pelo **CONTRATANTE** sem justa causa não o desobriga do "
        "pagamento dos honorários proporcionais ao trabalho prestado. O **CONTRATADO** "
        "poderá renunciar ao mandato mediante aviso prévio de 10 (dez) dias, continuando a "
        "representar o **CONTRATANTE** durante este período.")

    heading(doc, "CLÁUSULA 10ª — PROTEÇÃO DE DADOS (LGPD) E CONFIDENCIALIDADE")
    par(doc,
        "O **CONTRATANTE** consente expressamente (art. 7º, da Lei nº 13.709/18) com o "
        "tratamento de seus dados pessoais pelo **CONTRATADO**, exclusivamente para a "
        "finalidade de propositura e andamento da ação judicial. Ambas as partes se obrigam "
        "ao mais absoluto sigilo sobre as estratégias processuais e valores transacionados.")

    heading(doc, "CLÁUSULA 11ª — DAS DISPOSIÇÕES GERAIS E FORO")
    par(doc,
        "Para dirimir quaisquer controvérsias oriundas **DESTE CONTRATO**, as partes elegem "
        "o foro da Comarca de Parauapebas/PA, com renúncia expressa a qualquer outro.")

    par(doc, data_txt, align="right", antes=18, depois=6)
    assinatura(doc, d["cliente_nome"], f"Contratante — CPF n.º {d['cliente_cpf']}")
    assinatura(doc, ADV_ASSINA[0], ADV_ASSINA[1])
    out = os.path.join(saida_dir, f"CONTRATO DE HONORÁRIOS - {d['cliente_nome']}.docx")
    doc.save(out); return out

def _completar(d):
    if "cliente" not in d or not d["cliente"].strip():
        sys.exit("ERRO: campo 'cliente' (qualificação) é obrigatório.")
    d.setdefault("cliente_nome", d["cliente"].split(",")[0].strip())
    if not d.get("cliente_cpf"):
        m = re.search(r"\d{3}\.\d{3}\.\d{3}-\d{2}", d["cliente"])
        d["cliente_cpf"] = m.group(0) if m else "____________"
    d.setdefault("honorarios", "50% (CINQUENTA POR CENTO)")
    return d

def main():
    ap = argparse.ArgumentParser(description="Gera procuração/contrato de honorários (JEC) no timbrado Nascimento.")
    ap.add_argument("--dados", required=True, help="arquivo .json com os dados do caso")
    ap.add_argument("--saida-dir", default=".", help="pasta de saída (padrão: atual)")
    ap.add_argument("--data", default=None, help='texto da data (padrão: "Parauapebas/PA, <hoje>.")')
    ap.add_argument("--template", default=None, help="caminho do timbrado (padrão: assets/)")
    ap.add_argument("--so", choices=["procuracao","contrato","ambos"], default="ambos")
    a = ap.parse_args()
    with open(a.dados, encoding="utf-8") as fh:
        d = _completar(json.load(fh))
    template = a.template or _default_template()
    os.makedirs(a.saida_dir, exist_ok=True)
    if a.data:
        data_txt = a.data if a.data.strip().endswith(".") else a.data.strip()+"."
    else:
        h = datetime.date.today()
        data_txt = f"Parauapebas/PA, {h.day} de {MESES[h.month-1]} de {h.year}."
    feitos = []
    if a.so in ("procuracao","ambos"):
        feitos.append(gerar_procuracao(d, template, a.saida_dir, data_txt))
    if a.so in ("contrato","ambos"):
        for c in ("objeto","reus"):
            if not d.get(c): sys.exit(f"ERRO: campo '{c}' é obrigatório para o contrato.")
        feitos.append(gerar_contrato(d, template, a.saida_dir, data_txt))
    for f in feitos: print("OK ->", f)

if __name__ == "__main__":
    main()
